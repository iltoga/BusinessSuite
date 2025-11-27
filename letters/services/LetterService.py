import os
from datetime import date, datetime
from io import BytesIO
from uuid import uuid4

from django.conf import settings
from django.utils.timezone import now as datetime_now
from mailmerge import MailMerge

import core.utils.formatutils as formatutils
from customers.models import Customer


class LetterService:
    DATE_KEYS = {"doc_date", "birthdate", "passport_exp_date"}

    VISA_TYPES_SURAT_PERMOHONAN = {
        "voa": "VOA",
        "C1": "C1",
    }

    GENDER_TRANSLATIONS = {
        "m": "Laki-laki",
        "male": "Laki-laki",
        "f": "Perempuan",
        "female": "Perempuan",
        "laki-laki": "Laki-laki",
        "perempuan": "Perempuan",
    }

    def __init__(self, customer: Customer, template_name: str):
        self.customer = customer
        self.template_name = template_name
        self.generated_temp_files: list[str] = []

    def _format_date_value(self, value, format_type="slash"):
        if not value:
            return ""

        if isinstance(value, datetime):
            if format_type == "long":
                return formatutils.as_long_date_str(value)
            if format_type == "dash":
                return formatutils.as_date_dash_str(value)
            return formatutils.as_date_str(value)

        if isinstance(value, date):
            if format_type == "long":
                return formatutils.as_long_date_str(value)
            if format_type == "dash":
                return formatutils.as_date_dash_str(value)
            return formatutils.as_date_str(value)

        if isinstance(value, str):
            normalized_value = value.strip()
            if not normalized_value:
                return ""

            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
                try:
                    parsed = datetime.strptime(normalized_value, fmt)
                    if format_type == "long":
                        return formatutils.as_long_date_str(parsed)
                    if format_type == "dash":
                        return formatutils.as_date_dash_str(parsed)
                    return formatutils.as_date_str(parsed)
                except ValueError:
                    continue
            return normalized_value

        if format_type == "long":
            return formatutils.as_long_date_str(value)
        if format_type == "dash":
            return formatutils.as_date_dash_str(value)
        return formatutils.as_date_str(value)

    def _translate_gender(self, value):
        if not value:
            return ""
        normalized = str(value).strip()
        if not normalized:
            return ""
        translated = self.GENDER_TRANSLATIONS.get(normalized.lower())
        return translated if translated else normalized

    def generate_letter_data(self, extra_data=None):
        # Use shared formatting utility to ensure dates are formatted as '17 November 2025'
        cur_date = formatutils.as_long_date_str(datetime_now())

        nationality = self.customer.nationality.country if self.customer.nationality else ""

        data = {
            "doc_date": cur_date,
            "visa_type": self.VISA_TYPES_SURAT_PERMOHONAN.get("voa"),
            "name": self.customer.full_name,
            "gender": self._translate_gender(self.customer.gender or self.customer.get_gender_display()),
            "country": nationality,
            "birth_place": nationality,  # Default to country if not provided
            "birthdate": self._format_date_value(self.customer.birthdate, format_type="dash"),
            "passport_n": self.customer.passport_number or "",
            "passport_exp_date": self._format_date_value(self.customer.passport_expiration_date, format_type="dash"),
            "address_bali": self.customer.address_bali or "",
        }

        if extra_data:
            normalized = {}
            for key, value in extra_data.items():
                if key in self.DATE_KEYS:
                    if key == "doc_date":
                        normalized[key] = self._format_date_value(value, format_type="long")
                    elif key in {"birthdate", "passport_exp_date"}:
                        normalized[key] = self._format_date_value(value, format_type="dash")
                    else:
                        normalized[key] = self._format_date_value(value)
                else:
                    normalized[key] = value.strip() if isinstance(value, str) else value

            data.update({k: v for k, v in normalized.items() if v is not None})

            # Normalize visa_type to mapped display value if provided as key
            if data.get("visa_type"):
                data["visa_type"] = self.VISA_TYPES_SURAT_PERMOHONAN.get(data.get("visa_type"), data.get("visa_type"))

            # If address_bali is provided in extra_data, split into lines as well
            if "address_bali" in normalized:
                addr_lines = (normalized.get("address_bali") or "").splitlines()
                for i in range(4):
                    data[f"address_bali_line_{i+1}"] = addr_lines[i].strip() if i < len(addr_lines) else ""
                # replace address_bali with normalized newline-separated string
                data["address_bali"] = "\n".join([ln.strip() for ln in addr_lines if ln.strip()])

        data["gender"] = self._translate_gender(data.get("gender"))

        # If birth_place is empty, use country as fallback
        if not data.get("birth_place"):
            data["birth_place"] = data.get("country", "")

        # Split address lines into explicit merge fields so the template can render them on separate lines
        address_lines = (self.customer.address_bali or "").splitlines()
        # Normalize address_bali to use explicit newline separators and trim whitespace
        data["address_bali"] = "\n".join([ln.strip() for ln in address_lines if ln.strip()])
        for i in range(4):
            key = f"address_bali_line_{i+1}"
            data[key] = address_lines[i].strip() if i < len(address_lines) else ""

        return data

    def generate_letter_document(self, data):
        # Keep templates together with other DOCX templates under the 'reporting' folder
        template_path = os.path.join(settings.STATIC_SOURCE_ROOT, "reporting", self.template_name)

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Letter template '{self.template_name}' was not found at {template_path}.")

        with open(template_path, "rb") as template:
            doc = MailMerge(template)
            doc.merge(**data)

            # write to a buffer
            buf = BytesIO()
            doc.write(buf)

        # Post-process the document to ensure address line breaks are preserved as separate runs
        try:
            from docx import Document as DocxDocument

            buf.seek(0)
            docx_doc = DocxDocument(buf)

            # Determine address lines from data
            addr_lines = []
            if data.get("address_bali_line_1"):
                for i in range(1, 5):
                    line = data.get(f"address_bali_line_{i}")
                    if line:
                        addr_lines.append(line)
            elif data.get("address_bali"):
                addr_lines = [ln for ln in data.get("address_bali", "").splitlines() if ln.strip()]

            if addr_lines:
                # First look for the paragraph that contains the merged customer name so we can target the nearby 'Alamat'
                customer_name = data.get("name") or ""
                target_para = None
                if customer_name:
                    for idx, paragraph in enumerate(docx_doc.paragraphs):
                        if customer_name in (paragraph.text or ""):
                            # search in the next few paragraphs for 'Alamat'
                            for j in range(idx, min(idx + 10, len(docx_doc.paragraphs))):
                                p_text = (docx_doc.paragraphs[j].text or "").lower()
                                if "alamat" in p_text:
                                    target_para = docx_doc.paragraphs[j]
                                    break
                            if target_para:
                                break

                # fallback: find the first paragraph that contains 'Alamat' if customer_name search fails
                if not target_para:
                    for paragraph in docx_doc.paragraphs:
                        p_text = (paragraph.text or "").lower()
                        if "alamat" in p_text:
                            # If one of the address lines is already present in the paragraph, skip (template already contains lines)
                            if any(line in p_text for line in addr_lines):
                                target_para = None
                                break
                            target_para = paragraph
                            break

                if target_para:
                    paragraph = target_para
                    p_text = paragraph.text or ""
                    # find the first occurrence of the first address line in the paragraph text to preserve prefix
                    first_line = addr_lines[0]
                    idx = p_text.find(first_line)
                    if idx == -1:
                        prefix = p_text
                    else:
                        prefix = p_text[:idx]

                    # clear existing runs and set prefix
                    if paragraph.runs:
                        paragraph.runs[0].text = prefix
                        for r in paragraph.runs[1:]:
                            r.text = ""
                    else:
                        paragraph.add_run(prefix)

                    # append address lines as separate runs with breaks
                    for i, line in enumerate(addr_lines):
                        if not line:
                            continue
                        # Insert a break before the first appended line to separate from the label/prefix
                        if i == 0:
                            tmp_run = paragraph.add_run()
                            tmp_run.add_break()
                        r = paragraph.add_run(line)
                        if i < len(addr_lines) - 1:
                            r.add_break()

            # write final docx back to buffer
            final_buf = BytesIO()
            docx_doc.save(final_buf)
            final_buf.seek(0)
            final_bytes = final_buf.getvalue()
            self._persist_temp_file(final_bytes)
            return BytesIO(final_bytes)
        except Exception:
            # If anything goes wrong with docx post-processing, return the original buffer
            buf.seek(0)
            return buf

    def _persist_temp_file(self, content_bytes: bytes):
        tmp_folder = getattr(settings, "TMPFILES_FOLDER", "tmpfiles")
        media_root = getattr(settings, "MEDIA_ROOT", None)
        if media_root:
            tmp_dir = os.path.join(media_root, tmp_folder)
        else:
            tmp_dir = tmp_folder

        os.makedirs(tmp_dir, exist_ok=True)
        timestamp = datetime_now().strftime("%Y%m%d%H%M%S")
        filename = f"surat_permohonan_{timestamp}_{uuid4().hex}.docx"
        tmp_path = os.path.join(tmp_dir, filename)

        try:
            with open(tmp_path, "wb") as tmp_file:
                tmp_file.write(content_bytes)
            self.generated_temp_files.append(tmp_path)
        except OSError:
            pass
