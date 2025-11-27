"""
AI Invoice Parser Service
Uses AI vision to extract structured invoice data from documents.
Supports multimodal vision for PDF/images and text extraction for Excel/Word documents.
"""

import logging
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from typing import List, Optional, Union

from django.core.files.uploadedfile import UploadedFile

from core.services.ai_client import AIClient

logger = logging.getLogger(__name__)


@dataclass
class InvoiceLineItemData:
    """Structured data for a single invoice line item."""

    code: str
    description: str
    quantity: float
    unit_price: float
    amount: float
    notes: Optional[str] = None  # Person-specific details like names, numbers, etc.


@dataclass
class CustomerData:
    """Structured data for customer information."""

    full_name: str
    customer_type: str = "person"  # "person" or "company"
    company_name: Optional[str] = None
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    mobile_phone: Optional[str] = None
    npwp: Optional[str] = None
    address_bali: Optional[str] = None


@dataclass
class InvoiceData:
    """Structured data for invoice information."""

    invoice_no: str
    invoice_date: str  # ISO format: YYYY-MM-DD
    due_date: str  # ISO format: YYYY-MM-DD
    total_amount: float
    notes: Optional[str] = None
    payment_status: Optional[str] = None
    bank_details: Optional[dict] = None


@dataclass
class ParsedInvoiceResult:
    """Complete parsed invoice result."""

    customer: CustomerData
    invoice: InvoiceData
    line_items: List[InvoiceLineItemData]
    confidence_score: float  # 0-1 scale
    raw_response: dict


class AIInvoiceParser:
    """
    Service to parse invoices using AI vision.
    Supports direct PDF and image processing without OCR preprocessing.
    """

    # JSON Schema for structured output
    INVOICE_SCHEMA = {
        "type": "object",
        "properties": {
            "customer": {
                "type": "object",
                "properties": {
                    "full_name": {"type": "string"},
                    "customer_type": {"type": "string", "enum": ["person", "company"]},
                    "company_name": {"type": ["string", "null"]},
                    "first_name": {"type": ["string", "null"]},
                    "last_name": {"type": ["string", "null"]},
                    "email": {"type": ["string", "null"]},
                    "phone": {"type": ["string", "null"]},
                    "mobile_phone": {"type": ["string", "null"]},
                    "npwp": {"type": ["string", "null"]},
                    "address_bali": {"type": ["string", "null"]},
                },
                "required": [
                    "full_name",
                    "customer_type",
                    "company_name",
                    "first_name",
                    "last_name",
                    "email",
                    "phone",
                    "mobile_phone",
                    "npwp",
                    "address_bali",
                ],
                "additionalProperties": False,
            },
            "invoice": {
                "type": "object",
                "properties": {
                    "invoice_no": {"type": "string"},
                    "invoice_date": {"type": "string"},
                    "due_date": {"type": "string"},
                    "total_amount": {"type": "number"},
                    "notes": {"type": ["string", "null"]},
                    "payment_status": {"type": ["string", "null"]},
                    "bank_details": {
                        "type": ["object", "null"],
                        "properties": {
                            "bank_name": {"type": ["string", "null"]},
                            "beneficiary_name": {"type": ["string", "null"]},
                            "account_number": {"type": ["string", "null"]},
                            "branch": {"type": ["string", "null"]},
                            "address": {"type": ["string", "null"]},
                            "bic_swift": {"type": ["string", "null"]},
                            "email": {"type": ["string", "null"]},
                        },
                        "required": [
                            "bank_name",
                            "beneficiary_name",
                            "account_number",
                            "branch",
                            "address",
                            "bic_swift",
                            "email",
                        ],
                        "additionalProperties": False,
                    },
                },
                "required": [
                    "invoice_no",
                    "invoice_date",
                    "due_date",
                    "total_amount",
                    "notes",
                    "payment_status",
                    "bank_details",
                ],
                "additionalProperties": False,
            },
            "line_items": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "code": {"type": "string"},
                        "description": {"type": "string"},
                        "quantity": {"type": "number"},
                        "unit_price": {"type": "number"},
                        "amount": {"type": "number"},
                        "notes": {"type": ["string", "null"]},
                    },
                    "required": ["code", "description", "quantity", "unit_price", "amount", "notes"],
                    "additionalProperties": False,
                },
            },
            "confidence_score": {"type": "number", "minimum": 0, "maximum": 1},
        },
        "required": ["customer", "invoice", "line_items", "confidence_score"],
        "additionalProperties": False,
    }

    # Supported file types
    VISION_TYPES = ["pdf", "png", "jpg", "jpeg", "gif", "webp", "docx", "doc"]
    STRUCTURED_TYPES = ["xlsx", "xls"]

    # System prompt for invoice analysis
    SYSTEM_PROMPT = (
        "You are an expert at extracting structured data from invoice documents. "
        "Analyze the invoice image carefully and extract all relevant information."
    )

    def __init__(
        self,
        api_key: Optional[str] = None,
        model: Optional[str] = None,
        use_openrouter: Optional[bool] = None,
    ):
        """
        Initialize the parser.

        Args:
            api_key: API key (defaults to settings)
            model: Model to use (defaults to settings)
            use_openrouter: Whether to use OpenRouter (defaults to settings)
        """
        self.ai_client = AIClient(
            api_key=api_key,
            model=model,
            use_openrouter=use_openrouter,
        )
        logger.info(
            f"Initialized AI Invoice parser with {self.ai_client.provider_name} " f"(model: {self.ai_client.model})"
        )

    @property
    def model(self) -> str:
        """Get the model name for backward compatibility."""
        return self.ai_client.model

    @property
    def use_openrouter(self) -> bool:
        """Get whether using OpenRouter for backward compatibility."""
        return self.ai_client.use_openrouter

    @property
    def client(self):
        """Get the underlying OpenAI client for backward compatibility."""
        return self.ai_client.client

    def parse_invoice_file(
        self,
        file_content: Union[bytes, UploadedFile],
        filename: str = "",
        file_type: str = "",
    ) -> Optional[ParsedInvoiceResult]:
        """
        Parse invoice file using multimodal vision (PDF, images) or
        structured data extraction (Excel, Word).

        Args:
            file_content: File bytes or Django UploadedFile
            filename: Original filename for context
            file_type: File extension (pdf, png, jpg, xlsx, docx, etc.)

        Returns:
            ParsedInvoiceResult or None if parsing fails
        """
        try:
            # Read file bytes
            file_bytes, detected_filename = AIClient.read_file_bytes(file_content)
            filename = filename or detected_filename

            # Detect file type if not provided
            if not file_type:
                file_type = AIClient.get_file_extension(filename)

            logger.info(f"Parsing invoice file: {filename} " f"(type: {file_type}, model: {self.ai_client.model})")

            # For PDF, images, and Word documents, use multimodal vision
            if file_type in self.VISION_TYPES:
                return self._parse_with_vision(file_bytes, filename, file_type)

            # For Excel, extract text first then use LLM
            elif file_type in self.STRUCTURED_TYPES:
                return self._parse_structured_document(file_bytes, filename, file_type)

            else:
                logger.error(f"Unsupported file type: {file_type}")
                return None

        except Exception as e:
            logger.error(f"Error parsing invoice file: {str(e)}")
            return None

    def _parse_with_vision(
        self,
        file_bytes: bytes,
        filename: str,
        file_type: str,
    ) -> Optional[ParsedInvoiceResult]:
        """Parse PDF, DOCX, or image using vision capabilities."""
        try:
            # Convert document to image if needed
            image_bytes = self._convert_to_image(file_bytes, file_type)
            if image_bytes is None:
                return None

            # Build vision messages
            prompt = self._build_vision_prompt(filename)
            messages = self.ai_client.build_vision_message(
                prompt=prompt,
                image_bytes=image_bytes,
                filename=f"{filename}.png",  # Force PNG MIME type
                system_prompt=self.SYSTEM_PROMPT,
            )

            logger.info(f"Sending invoice image to {self.ai_client.provider_name} vision API")

            # Call API with structured output
            parsed_data = self.ai_client.chat_completion_json(
                messages=messages,
                json_schema=self.INVOICE_SCHEMA,
                schema_name="invoice_data",
            )

            logger.info("Successfully parsed invoice data from vision API")

            return self._convert_to_result(parsed_data)

        except Exception as e:
            logger.error(f"Error in vision parsing: {str(e)}")
            return None

    def _convert_to_image(self, file_bytes: bytes, file_type: str) -> Optional[bytes]:
        """Convert document to image bytes for vision API."""
        try:
            if file_type == "pdf":
                from pdf2image import convert_from_bytes

                images = convert_from_bytes(file_bytes, dpi=200, fmt="png")
                img_byte_arr = BytesIO()
                images[0].save(img_byte_arr, format="PNG")
                return img_byte_arr.getvalue()

            elif file_type in ["docx", "doc"]:
                from docx import Document
                from PIL import Image, ImageDraw, ImageFont

                # Extract text and tables from DOCX
                doc = Document(BytesIO(file_bytes))
                lines = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        lines.append(para.text)

                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            lines.append(row_text)

                # Create image from text
                return self._text_to_image(lines)

            else:
                # Already an image
                return file_bytes

        except Exception as e:
            logger.error(f"Error converting document to image: {str(e)}")
            return None

    def _text_to_image(self, lines: List[str]) -> bytes:
        """Convert text lines to an image."""
        from PIL import Image, ImageDraw, ImageFont

        img_width = 1200
        font_size = 20
        line_height = 30
        padding = 40

        try:
            font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", font_size)
        except Exception:
            font = ImageFont.load_default()

        # Calculate image height
        img_height = max(800, len(lines) * line_height + padding * 2)

        img = Image.new("RGB", (img_width, img_height), color="white")
        draw = ImageDraw.Draw(img)

        y_position = padding
        for line in lines:
            # Wrap long lines
            if len(line) > 80:
                words = line.split()
                current_line = ""
                for word in words:
                    if len(current_line + " " + word) <= 80:
                        current_line += " " + word if current_line else word
                    else:
                        draw.text((padding, y_position), current_line, fill="black", font=font)
                        y_position += line_height
                        current_line = word
                if current_line:
                    draw.text((padding, y_position), current_line, fill="black", font=font)
                    y_position += line_height
            else:
                draw.text((padding, y_position), line, fill="black", font=font)
                y_position += line_height

        img_byte_arr = BytesIO()
        img.save(img_byte_arr, format="PNG")
        return img_byte_arr.getvalue()

    def _parse_structured_document(
        self,
        file_bytes: bytes,
        filename: str,
        file_type: str,
    ) -> Optional[ParsedInvoiceResult]:
        """Parse Excel or Word documents by extracting text then using LLM."""
        try:
            # Extract text
            if file_type in ["xlsx", "xls"]:
                text = self._extract_excel_text(file_bytes)
            elif file_type in ["docx", "doc"]:
                text = self._extract_docx_text(file_bytes)
            else:
                logger.error(f"Unsupported structured document type: {file_type}")
                return None

            # Build messages for text parsing
            prompt = self._build_text_prompt(text, filename)
            messages = [
                {"role": "system", "content": self.SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ]

            logger.info(f"Sending extracted text to {self.ai_client.provider_name}")

            # Call API with structured output
            parsed_data = self.ai_client.chat_completion_json(
                messages=messages,
                json_schema=self.INVOICE_SCHEMA,
                schema_name="invoice_data",
            )

            logger.info("Successfully parsed invoice data from structured document")

            return self._convert_to_result(parsed_data)

        except Exception as e:
            logger.error(f"Error parsing structured document: {str(e)}")
            return None

    def _extract_excel_text(self, file_bytes: bytes) -> str:
        """Extract text from Excel file."""
        from datetime import datetime as dt

        import openpyxl

        workbook = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
        sheet = workbook.active
        text_parts = []

        for row in sheet.iter_rows(values_only=True):
            formatted_cells = []
            for cell in row:
                if cell is None:
                    formatted_cells.append("")
                elif isinstance(cell, dt):
                    formatted_cells.append(cell.strftime("%Y-%m-%d"))
                elif isinstance(cell, (int, float)):
                    formatted_cells.append(str(cell))
                else:
                    formatted_cells.append(str(cell))

            row_text = " | ".join(formatted_cells)
            if row_text.strip():
                text_parts.append(row_text)

        return "\n".join(text_parts)

    def _extract_docx_text(self, file_bytes: bytes) -> str:
        """Extract text from Word document."""
        from docx import Document

        doc = Document(BytesIO(file_bytes))
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    def _build_vision_prompt(self, filename: str) -> str:
        """Build prompt for vision API (image/PDF analysis)."""
        return f"""Analyze this invoice document (file: {filename}) and extract all structured data.

IMPORTANT EXTRACTION RULES:

1. Extract ALL line items/services listed
2. Convert dates to YYYY-MM-DD format (e.g., "10/29/2025" → "2025-10-29", "29 Oct 2025" → "2025-10-29")
3. Remove currency symbols and formatting from amounts (e.g., "Rp 16,250,000" → 16250000.00)
4. For invoice_no, extract only numeric part (e.g., "INV-202634" → "202634")
5. Split customer full_name into first_name and last_name
6. If due_date not specified, use same as invoice_date
7. confidence_score: 0.9+ if clear, 0.5-0.8 if partially unclear, <0.5 if very uncertain
8. For missing optional fields, use null

9. **CRITICAL - CUSTOMER TYPE DETECTION:**
   Analyze the "Bill To:" section to determine customer_type:
   - If it contains ONLY a person's first and last name → customer_type: "person"
   - If it contains ONLY a company name (often starting with "PT.", "CV.", "Ltd.", "Inc.", etc.) → customer_type: "company"
   - If BOTH company name AND person name are present → customer_type: "person" (populate both company_name and first_name/last_name)

   For PERSON customers:
   - Extract first_name and last_name
   - Extract company_name (if present)

   For COMPANY customers:
   - Extract company_name (the business name)

   Additional fields to try to extract from "Bill To:" section:
   - phone or mobile_phone (telephone number)
   - npwp (Indonesian tax ID, usually 15 digits with dots/dashes)
   - address_bali (address)

10. **CRITICAL - LINE ITEMS WITH MULTIPLE PEOPLE:**
   If a line item description mentions MULTIPLE people or has quantity > 1:
   - CREATE SEPARATE LINE ITEMS for EACH person
   - Set quantity=1 for each separate line item
   - Put ONLY ONE person's name in the 'notes' field for each line item
   - Use unit_price (not total) for amount

   Example Input: "XVOA | Extension visa on arrival for Marco Polo, Ms. Beatrice Manzoni | QTY: 2 | Price: Rp 900,000 | Total: Rp 1,800,000"

   Correct Output (2 separate line items):
   Line Item 1: {{
     "code": "XVOA",
     "description": "Extension visa on arrival for",
     "quantity": 1,
     "unit_price": 900000.00,
     "amount": 900000.00,
     "notes": "Marco Polo"
   }}
   Line Item 2: {{
     "code": "XVOA",
     "description": "Extension visa on arrival for",
     "quantity": 1,
     "unit_price": 900000.00,
     "amount": 900000.00,
     "notes": "Beatrice Manzoni"
   }}

   WRONG (do NOT do this): One line item with quantity=2 and both names in notes

11. For line items with only ONE person:
    - Extract person name/details to 'notes' field
    - Keep generic service description in 'description'
    - Example: "Visa for John Smith" → description: "Visa for", notes: "John Smith"

Look carefully at the document for:
- Customer type (person vs company) in "Bill To:" section
- Company name (if present, often starts with PT., CV., etc.)
- Person name (first and last name)
- Telephone number
- NPWP (tax ID)
- Address in Bali
- Invoice number and dates
- Line items with codes, descriptions, quantities, prices
- Multiple people mentioned in a single line item (split into separate items!)
- Total amount
- Bank details and payment info
"""

    def _build_text_prompt(self, text: str, filename: str) -> str:
        """Build prompt for text-based parsing (Excel/Word extracted text)."""
        return f"""Extract structured invoice data from the following text (file: {filename}).

INVOICE TEXT:
---
{text}
---

{self._build_vision_prompt(filename)}
"""

    def _convert_to_result(self, parsed_data: dict) -> ParsedInvoiceResult:
        """Convert parsed JSON data to structured result objects."""
        # Extract customer data
        customer_dict = parsed_data.get("customer", {})
        customer = CustomerData(
            full_name=customer_dict.get("full_name", ""),
            customer_type=customer_dict.get("customer_type", "person"),
            company_name=customer_dict.get("company_name"),
            first_name=customer_dict.get("first_name"),
            last_name=customer_dict.get("last_name"),
            email=customer_dict.get("email"),
            phone=customer_dict.get("phone"),
            mobile_phone=customer_dict.get("mobile_phone"),
            npwp=customer_dict.get("npwp"),
            address_bali=customer_dict.get("address_bali"),
        )

        # Extract invoice data
        invoice_dict = parsed_data.get("invoice", {})
        invoice = InvoiceData(
            invoice_no=str(invoice_dict.get("invoice_no", "")),
            invoice_date=invoice_dict.get("invoice_date", ""),
            due_date=invoice_dict.get("due_date", ""),
            total_amount=float(invoice_dict.get("total_amount", 0)),
            notes=invoice_dict.get("notes"),
            payment_status=invoice_dict.get("payment_status"),
            bank_details=invoice_dict.get("bank_details"),
        )

        # Extract line items
        line_items = []
        for item_dict in parsed_data.get("line_items", []):
            line_item = InvoiceLineItemData(
                code=item_dict.get("code", ""),
                description=item_dict.get("description", ""),
                quantity=float(item_dict.get("quantity", 1)),
                unit_price=float(item_dict.get("unit_price", 0)),
                amount=float(item_dict.get("amount", 0)),
                notes=item_dict.get("notes"),
            )
            line_items.append(line_item)

        confidence_score = float(parsed_data.get("confidence_score", 0.5))

        return ParsedInvoiceResult(
            customer=customer,
            invoice=invoice,
            line_items=line_items,
            confidence_score=confidence_score,
            raw_response=parsed_data,
        )

    def validate_parsed_data(self, result: ParsedInvoiceResult) -> tuple[bool, list[str]]:
        """
        Validate parsed data for completeness and consistency.

        Returns:
            (is_valid, list_of_errors)
        """
        errors = []

        # Validate customer
        if not (result.customer.full_name or result.customer.company_name):
            errors.append("Customer name or company name is missing")

        # Validate invoice
        if not result.invoice.invoice_no:
            errors.append("Invoice number is missing")

        if not result.invoice.invoice_date:
            errors.append("Invoice date is missing")
        else:
            try:
                datetime.strptime(result.invoice.invoice_date, "%Y-%m-%d")
            except ValueError:
                errors.append(f"Invalid invoice date format: {result.invoice.invoice_date}")

        if not result.invoice.due_date:
            errors.append("Due date is missing")
        else:
            try:
                datetime.strptime(result.invoice.due_date, "%Y-%m-%d")
            except ValueError:
                errors.append(f"Invalid due date format: {result.invoice.due_date}")

        # Validate line items
        if not result.line_items:
            errors.append("No line items found")
        else:
            for i, item in enumerate(result.line_items):
                if not item.description:
                    errors.append(f"Line item {i+1}: description is missing")
                if item.amount < 0:
                    errors.append(f"Line item {i+1}: amount cannot be negative")

        # Validate total matches sum of line items
        if result.line_items:
            line_items_total = sum(item.amount for item in result.line_items)
            if abs(line_items_total - result.invoice.total_amount) > 0.01:
                errors.append(
                    f"Total amount mismatch: invoice total {result.invoice.total_amount} "
                    f"!= line items total {line_items_total}"
                )

        # Check confidence score
        if result.confidence_score < 0.5:
            errors.append(f"Low confidence score: {result.confidence_score:.2f}")

        return len(errors) == 0, errors

    def generate_product_details(self, code: str, description: str) -> dict:
        """
        Use LLM to generate a meaningful product name and sanitize description.

        Args:
            code: Product code from invoice line item
            description: Original description from invoice (may contain person-specific details)

        Returns:
            dict with 'name' and 'description' keys
        """
        try:
            prompt = f"""You are analyzing a product from an invoice to create a generic product catalog entry.

Product Code: {code or 'N/A'}
Original Description: {description}

Generate:
1. A short, meaningful product name (3-7 words) that combines the code and describes what the product/service is
2. A sanitized description by ONLY removing person-specific information (names, personal details, specific ID numbers) from the original description. DO NOT add extra information or rewrite it completely - just clean it up.

The description should be a cleaned version of the original, reusable for any customer.

Return ONLY a JSON object with this exact structure:
{{
    "name": "short meaningful product name",
    "description": "sanitized description without personal info"
}}"""

            messages = [
                {
                    "role": "system",
                    "content": "You are a product catalog assistant. Generate concise, generic product information suitable for a catalog.",
                },
                {"role": "user", "content": prompt},
            ]

            result = self.ai_client.chat_completion_simple_json(
                messages=messages,
                temperature=0.3,
            )

            # Ensure result is a dict, not a list
            if isinstance(result, list):
                logger.warning(f"LLM returned a list instead of dict: {result}")
                if result and isinstance(result[0], dict):
                    result = result[0]
                else:
                    raise ValueError("LLM returned invalid format (list)")

            if not isinstance(result, dict) or "name" not in result or "description" not in result:
                raise ValueError(f"LLM returned invalid structure: {result}")

            logger.info(f"Generated product details: {result.get('name', 'N/A')}")
            return result

        except Exception as e:
            logger.error(f"Error generating product details: {str(e)}", exc_info=True)
            # Fallback to basic generation
            fallback_name = f"{code} - {description.split()[0:3]}" if code else description[:50]
            return {
                "name": fallback_name,
                "description": description,
            }


# Backward compatibility alias
LLMInvoiceParser = AIInvoiceParser
