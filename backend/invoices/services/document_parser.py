"""
Document Parser Service
Extracts text content from PDF, Excel, and Word documents for invoice processing.
"""

import logging
import os
from typing import Callable, Optional

from core.services.logger_service import Logger

logger = Logger.get_logger(__name__)

import openpyxl
import pytesseract
from docx import Document
from pdf2image import convert_from_path
from PIL import Image


class DocumentParser:
    """
    Service to extract text from various document formats.
    """

    @staticmethod
    def extract_text_from_file(
        file_path: str,
        progress_callback: Optional[Callable[[int], None]] = None,
    ) -> Optional[str]:
        """
        Extract text from a document based on its file extension.

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text as string or None if extraction fails
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None

        file_extension = os.path.splitext(file_path)[1].lower()

        try:
            if file_extension == ".pdf":
                return DocumentParser._extract_from_pdf(file_path, progress_callback=progress_callback)
            elif file_extension in [".xlsx", ".xls"]:
                return DocumentParser._extract_from_excel(file_path, progress_callback=progress_callback)
            elif file_extension in [".docx", ".doc"]:
                return DocumentParser._extract_from_word(file_path, progress_callback=progress_callback)
            elif file_extension in [".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"]:
                return DocumentParser._extract_from_image(file_path, progress_callback=progress_callback)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None

    @staticmethod
    def _extract_from_pdf(
        file_path: str,
        progress_callback: Optional[Callable[[int], None]] = None,
    ) -> str:
        """
        Extract text from PDF using OCR.
        Uses pdf2image + pytesseract for text extraction.
        """
        try:
            # Convert PDF to images
            images = convert_from_path(file_path, dpi=300)

            extracted_text = []
            total_pages = len(images)
            for i, image in enumerate(images):
                logger.info(f"Processing page {i + 1}/{total_pages} from PDF")

                # Use OCR to extract text
                text = pytesseract.image_to_string(image, lang="eng")
                extracted_text.append(text)

                if progress_callback and total_pages > 0:
                    # Smoothly progress from 40% to 90% while pages are being OCRed.
                    progress = 40 + int(((i + 1) / total_pages) * 50)
                    progress_callback(min(progress, 90))

            full_text = "\n\n--- PAGE BREAK ---\n\n".join(extracted_text)
            logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
            return full_text

        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
            raise

    @staticmethod
    def _extract_from_excel(
        file_path: str,
        progress_callback: Optional[Callable[[int], None]] = None,
    ) -> str:
        """
        Extract text from Excel files.
        Reads all sheets and cells, formats them as structured text.
        """
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            if progress_callback:
                progress_callback(55)
            extracted_text = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                extracted_text.append(f"=== SHEET: {sheet_name} ===\n")

                # Extract cell values row by row
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    # Only add non-empty rows
                    if any(val.strip() for val in row_values):
                        extracted_text.append(" | ".join(row_values))

                extracted_text.append("\n")

            full_text = "\n".join(extracted_text)
            if progress_callback:
                progress_callback(90)
            logger.info(f"Successfully extracted {len(full_text)} characters from Excel")
            return full_text

        except Exception as e:
            logger.error(f"Error extracting from Excel: {str(e)}")
            raise

    @staticmethod
    def _extract_from_word(
        file_path: str,
        progress_callback: Optional[Callable[[int], None]] = None,
    ) -> str:
        """
        Extract text from Word documents.
        Reads all paragraphs and tables.
        """
        try:
            doc = Document(file_path)
            if progress_callback:
                progress_callback(55)
            extracted_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                extracted_text.append("\n=== TABLE ===")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        extracted_text.append(row_text)
                extracted_text.append("=== END TABLE ===\n")

            full_text = "\n".join(extracted_text)
            if progress_callback:
                progress_callback(90)
            logger.info(f"Successfully extracted {len(full_text)} characters from Word document")
            return full_text

        except Exception as e:
            logger.error(f"Error extracting from Word: {str(e)}")
            raise

    @staticmethod
    def _extract_from_image(
        file_path: str,
        progress_callback: Optional[Callable[[int], None]] = None,
    ) -> str:
        """Extract text from image files using OCR."""
        try:
            if progress_callback:
                progress_callback(55)
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang="eng")
            if progress_callback:
                progress_callback(90)
            logger.info(f"Successfully extracted {len(text)} characters from image")
            return text
        except Exception as e:
            logger.error(f"Error extracting from image: {str(e)}")
            raise

    @staticmethod
    def extract_text_from_uploaded_file(uploaded_file) -> Optional[str]:
        """
        Extract text from Django UploadedFile object.
        Saves temporarily, extracts, then cleans up.

        Args:
            uploaded_file: Django UploadedFile instance

        Returns:
            Extracted text as string or None if extraction fails
        """
        import tempfile

        try:
            # Create temporary file with proper extension
            file_extension = os.path.splitext(uploaded_file.name)[1]
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                # Write uploaded file to temp location
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_path = temp_file.name

            # Extract text
            extracted_text = DocumentParser.extract_text_from_file(temp_path)

            # Clean up temp file
            os.unlink(temp_path)

            return extracted_text

        except Exception as e:
            logger.error(f"Error processing uploaded file {uploaded_file.name}: {str(e)}")
            # Clean up temp file if it exists
            try:
                if "temp_path" in locals():
                    os.unlink(temp_path)
            except:
                pass
            return None
