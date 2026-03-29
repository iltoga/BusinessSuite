"""Service for building product price lists and related output artifacts."""

from __future__ import annotations

from collections import OrderedDict
from dataclasses import dataclass
from datetime import datetime
from decimal import Decimal
from io import BytesIO

from django.conf import settings
from django.utils import timezone
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from products.models import Product


@dataclass(frozen=True)
class PriceListProductRow:
    code: str
    name: str
    retail_price: Decimal | None
    currency: str


@dataclass(frozen=True)
class PriceListCategorySection:
    product_type: str
    product_type_label: str
    category_name: str
    products: tuple[PriceListProductRow, ...]


class ProductPriceListService:
    """Build a public-facing grouped product price list document.

    The generated document intentionally exposes only public retail prices.
    Base prices and other internal data must never be rendered.
    """

    title = "PRICE LIST"
    subtitle = "Retail prices only"
    accent_color = "8B6B2E"
    accent_fill = "E8D7B0"
    muted_fill = "F7F1E4"
    border_color = "D8C49A"
    company_name = getattr(settings, "PRICE_LIST_COMPANY_NAME", "PT. BALI SHANKARA KONSULTAN")
    company_address = getattr(
        settings,
        "PRICE_LIST_COMPANY_ADDRESS",
        "Jl. Raya Nusa Dua Selatan Ruko F, Kutuh, Kuta Selatan, Badung - Bali, Indonesia 80361",
    )
    company_contact = getattr(
        settings,
        "PRICE_LIST_COMPANY_CONTACT",
        "Phone +62 821 4720 9978 · +62 822 3629 6258 · Email info@revisbali.com",
    )
    company_website = getattr(settings, "PRICE_LIST_COMPANY_WEBSITE", "www.revisbali.com")

    def build_sections(self) -> list[PriceListCategorySection]:
        queryset = (
            Product.objects.select_related("product_category")
            .filter(deprecated=False)
            .order_by("product_category__product_type", "product_category__name", "code", "name")
        )

        grouped: OrderedDict[tuple[str, str], list[PriceListProductRow]] = OrderedDict()
        for product in queryset.iterator():
            category = product.product_category
            if not category:
                continue
            key = (category.product_type, category.name)
            grouped.setdefault(key, []).append(
                PriceListProductRow(
                    code=str(product.code or "").strip() or "—",
                    name=str(product.name or "").strip() or "Unnamed service",
                    retail_price=product.retail_price,
                    currency=str(product.currency or "IDR").strip().upper() or "IDR",
                )
            )

        sections: list[PriceListCategorySection] = []
        for (product_type, category_name), products in grouped.items():
            sections.append(
                PriceListCategorySection(
                    product_type=product_type,
                    product_type_label=self._product_type_label(product_type),
                    category_name=category_name,
                    products=tuple(products),
                )
            )
        return sections

    def generate_docx_buffer(
        self,
        sections: list[PriceListCategorySection],
        *,
        generated_at: datetime | None = None,
    ) -> tuple[BytesIO, dict[str, int | str]]:
        document = Document()
        generated_at = generated_at or timezone.localtime()

        self._configure_document(document)
        self._add_header(document, generated_at)

        current_product_type: str | None = None
        for index, section in enumerate(sections):
            if current_product_type != section.product_type:
                if index > 0:
                    document.add_page_break()
                self._add_product_type_intro(document, section.product_type_label)
                current_product_type = section.product_type

            self._add_category_section(document, section)

        if not sections:
            empty = document.add_paragraph()
            empty.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = empty.add_run("No active products are currently available for the public price list.")
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(120, 120, 120)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        summary = {
            "total_categories": len(sections),
            "total_products": sum(len(section.products) for section in sections),
            "generated_at": generated_at.isoformat(),
        }
        return buffer, summary

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        style = document.styles["Normal"]
        style.font.name = "Aptos"
        style.font.size = Pt(10)

        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.space_before = Pt(4)
        footer_para.space_after = Pt(0)
        footer_run = footer_para.add_run(
            f"{self.company_name} · {self.company_address}\n{self.company_contact} · {self.company_website}"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(112, 112, 112)

    def _add_header(self, document: Document, generated_at: datetime) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(4)
        title_run = title.add_run(self.title)
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = self._rgb_from_hex(self.accent_color)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(10)
        subtitle_run = subtitle.add_run(f"{self.subtitle} · Updated {generated_at.strftime('%d %B %Y')}")
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(11)

        note = document.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.paragraph_format.space_after = Pt(12)
        note_run = note.add_run(
            "Prepared for customer-facing use. Internal/base prices and non-public operational details are intentionally omitted."
        )
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(103, 103, 103)

    def _add_product_type_intro(self, document: Document, product_type_label: str) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after = Pt(6)
        run = heading.add_run(product_type_label.upper())
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = self._rgb_from_hex(self.accent_color)

        divider = document.add_paragraph()
        divider.paragraph_format.space_after = Pt(6)
        divider_run = divider.add_run("—" * 48)
        divider_run.font.color.rgb = self._rgb_from_hex(self.border_color)

    def _add_category_section(self, document: Document, section: PriceListCategorySection) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(4)
        run = heading.add_run(section.category_name)
        run.bold = True
        run.font.size = Pt(14)

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.autofit = False

        widths = [Cm(3.0), Cm(11.2), Cm(4.3)]
        headers = ["Code", "Service", "Retail price"]
        header_cells = table.rows[0].cells
        for idx, cell in enumerate(header_cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._set_cell_fill(cell, self.accent_fill)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx < 2 else WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(headers[idx])
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = self._rgb_from_hex(self.accent_color)

        for row_index, product in enumerate(section.products):
            row = table.add_row().cells
            for idx, cell in enumerate(row):
                cell.width = widths[idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if row_index % 2 == 0:
                    self._set_cell_fill(cell, self.muted_fill)

            code_para = row[0].paragraphs[0]
            code_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            code_run = code_para.add_run(product.code)
            code_run.bold = True
            code_run.font.size = Pt(10)

            service_para = row[1].paragraphs[0]
            service_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            service_run = service_para.add_run(product.name)
            service_run.font.size = Pt(10)

            price_para = row[2].paragraphs[0]
            price_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            price_run = price_para.add_run(self._format_currency(product.retail_price, product.currency))
            price_run.bold = True
            price_run.font.size = Pt(10)
            price_run.font.color.rgb = self._rgb_from_hex(self.accent_color)

        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)

    def _format_currency(self, value: Decimal | None, currency: str) -> str:
        if value is None:
            return "—"
        decimals = 0 if value == value.quantize(Decimal("1")) else 2
        formatted = f"{value:,.{decimals}f}"
        formatted = formatted.replace(",", "_").replace(".", ",").replace("_", ".")
        return f"{currency} {formatted}"

    def _product_type_label(self, product_type: str) -> str:
        normalized = str(product_type or "other").strip().lower()
        if normalized == "visa":
            return "Visa Services"
        if normalized == "other":
            return "Other Services"
        return normalized.replace("_", " ").title()

    def _set_cell_fill(self, cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def _rgb_from_hex(self, value: str) -> RGBColor:
        cleaned = value.strip().lstrip("#")
        return RGBColor.from_string(cleaned)
