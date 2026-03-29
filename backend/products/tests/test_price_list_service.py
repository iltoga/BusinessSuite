"""Tests for the product price list service."""

from decimal import Decimal

from django.test import TestCase
from docx import Document
from products.models import Product, ProductCategory
from products.services.price_list_service import ProductPriceListService


class ProductPriceListServiceTests(TestCase):
    def test_generate_docx_uses_only_retail_prices(self):
        visa_category = ProductCategory.objects.create(name="Single Entry Visa", product_type="visa")
        other_category = ProductCategory.objects.create(name="Company Services", product_type="other")

        Product.objects.create(
            product_category=visa_category,
            code="PR-1",
            name="Priority Visa",
            base_price=Decimal("111111.00"),
            retail_price=Decimal("999999.00"),
            currency="IDR",
        )
        Product.objects.create(
            product_category=other_category,
            code="PR-2",
            name="Company Setup",
            base_price=Decimal("222222.00"),
            retail_price=Decimal("333333.00"),
            currency="USD",
        )

        service = ProductPriceListService()
        sections = service.build_sections()

        self.assertEqual(len(sections), 2)
        sections_by_category = {section.category_name: section for section in sections}
        self.assertEqual(
            sections_by_category["Single Entry Visa"].products[0].retail_price,
            Decimal("999999.00"),
        )
        self.assertEqual(
            sections_by_category["Company Services"].products[0].retail_price,
            Decimal("333333.00"),
        )

        buffer, summary = service.generate_docx_buffer(sections)
        document = Document(buffer)
        paragraphs_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        full_text = f"{paragraphs_text}\n{table_text}"

        self.assertEqual(summary["total_categories"], 2)
        self.assertEqual(summary["total_products"], 2)
        self.assertIn("PRICE LIST", full_text)
        self.assertIn("Priority Visa", full_text)
        self.assertIn("IDR 999.999", full_text)
        self.assertIn("USD 333.333", full_text)
        self.assertNotIn("111.111", full_text)
        self.assertNotIn("222.222", full_text)
